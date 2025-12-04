"""
document_processor.py - Processamento de Documentos e Chunking
Extrai texto de PDF, Word e imagens escaneadas, e divide em chunks
"""

import os
import hashlib
from typing import List, Dict, Any, Optional
from io import BytesIO

# PDF Processing
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

# Word Processing
try:
    from docx import Document
except ImportError:
    Document = None

# OCR for Scanned Documents
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_bytes
except ImportError:
    pytesseract = None
    Image = None
    convert_from_bytes = None


class DocumentProcessor:
    """Processador de documentos com suporte a múltiplos formatos"""
    
    # Mapeamento de MIME types para extensões
    MIME_TO_EXT = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'image/png': 'png',
        'image/jpeg': 'jpg',
        'image/jpg': 'jpg',
        'image/tiff': 'tiff',
        'image/tif': 'tif',
    }
    
    # Extensões suportadas
    SUPPORTED_EXTENSIONS = ['pdf', 'docx', 'png', 'jpg', 'jpeg', 'tiff', 'tif']
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Inicializa o processador de documentos
        
        Args:
            chunk_size: Tamanho dos chunks em caracteres
            chunk_overlap: Sobreposição entre chunks em caracteres
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Verifica dependências
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Verifica se as dependências necessárias estão instaladas"""
        if PdfReader is None:
            print("[doc_processor] AVISO: PyPDF2 não instalado - processamento de PDF desabilitado")
        
        if Document is None:
            print("[doc_processor] AVISO: python-docx não instalado - processamento de Word desabilitado")
        
        if pytesseract is None or Image is None:
            print("[doc_processor] AVISO: pytesseract/PIL não instalados - OCR desabilitado")
    
    @staticmethod
    def calculate_hash(content: bytes) -> str:
        """
        Calcula hash SHA-256 do conteúdo
        
        Args:
            content: Conteúdo em bytes
            
        Returns:
            Hash hexadecimal
        """
        return hashlib.sha256(content).hexdigest()
    
    @staticmethod
    def is_supported_file(filename: str, mime_type: str = None) -> bool:
        """
        Verifica se o arquivo é suportado
        
        Args:
            filename: Nome do arquivo
            mime_type: MIME type do arquivo (opcional)
            
        Returns:
            True se suportado
        """
        # Verifica por extensão
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext in DocumentProcessor.SUPPORTED_EXTENSIONS:
            return True
        
        # Verifica por MIME type
        if mime_type and mime_type in DocumentProcessor.MIME_TO_EXT:
            return True
        
        return False
    
    def extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extrai texto de arquivo PDF
        
        Args:
            content: Conteúdo do PDF em bytes
            
        Returns:
            Texto extraído
        """
        if PdfReader is None:
            raise RuntimeError("PyPDF2 não está instalado")
        
        try:
            pdf_file = BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n\n".join(text_parts)
            
            # Se não conseguiu extrair texto, tenta OCR
            if not full_text.strip() and convert_from_bytes is not None:
                print("[doc_processor] PDF sem texto extraível - tentando OCR...")
                return self._ocr_from_pdf(content)
            
            return full_text
            
        except Exception as e:
            raise RuntimeError(f"Erro ao processar PDF: {str(e)}")
    
    def extract_text_from_docx(self, content: bytes) -> str:
        """
        Extrai texto de arquivo Word (.docx)
        
        Args:
            content: Conteúdo do Word em bytes
            
        Returns:
            Texto extraído
        """
        if Document is None:
            raise RuntimeError("python-docx não está instalado")
        
        try:
            docx_file = BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extrai texto de parágrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extrai texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"Erro ao processar Word: {str(e)}")
    
    def extract_text_from_image(self, content: bytes) -> str:
        """
        Extrai texto de imagem usando OCR
        
        Args:
            content: Conteúdo da imagem em bytes
            
        Returns:
            Texto extraído via OCR
        """
        if pytesseract is None or Image is None:
            raise RuntimeError("pytesseract/PIL não estão instalados")
        
        try:
            image = Image.open(BytesIO(content))
            
            # Configura OCR para português e inglês
            custom_config = r'--oem 3 --psm 6 -l por+eng'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            return text
            
        except Exception as e:
            raise RuntimeError(f"Erro ao processar imagem com OCR: {str(e)}")
    
    def _ocr_from_pdf(self, content: bytes) -> str:
        """
        Extrai texto de PDF usando OCR (para PDFs escaneados)
        
        Args:
            content: Conteúdo do PDF em bytes
            
        Returns:
            Texto extraído via OCR
        """
        if convert_from_bytes is None or pytesseract is None:
            raise RuntimeError("pdf2image/pytesseract não estão instalados")
        
        try:
            # Converte PDF para imagens
            images = convert_from_bytes(content)
            
            text_parts = []
            custom_config = r'--oem 3 --psm 6 -l por+eng'
            
            for i, image in enumerate(images):
                print(f"[doc_processor] Processando página {i+1}/{len(images)} com OCR...")
                text = pytesseract.image_to_string(image, config=custom_config)
                if text.strip():
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"Erro ao processar PDF com OCR: {str(e)}")
    
    def extract_text(self, content: bytes, filename: str, mime_type: str = None) -> str:
        """
        Extrai texto do documento baseado no tipo
        
        Args:
            content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo
            mime_type: MIME type do arquivo
            
        Returns:
            Texto extraído
        """
        # Determina o tipo do arquivo
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        
        if not ext and mime_type:
            ext = self.MIME_TO_EXT.get(mime_type, '')
        
        # Processa baseado no tipo
        if ext == 'pdf' or mime_type == 'application/pdf':
            return self.extract_text_from_pdf(content)
        
        elif ext == 'docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self.extract_text_from_docx(content)
        
        elif ext in ['png', 'jpg', 'jpeg', 'tiff', 'tif'] or (mime_type and mime_type.startswith('image/')):
            return self.extract_text_from_image(content)
        
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {ext or mime_type}")
    
    def create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        Divide o texto em chunks com sobreposição
        
        Args:
            text: Texto completo para dividir
            
        Returns:
            Lista de dicionários com informações dos chunks
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        text_length = len(text)
        start = 0
        chunk_index = 0
        
        while start < text_length:
            # Define o fim do chunk
            end = start + self.chunk_size
            
            # Se não é o último chunk, tenta quebrar em espaço ou pontuação
            if end < text_length:
                # Procura por quebra natural (ponto, nova linha, espaço)
                search_start = end
                search_end = min(end + 100, text_length)
                
                # Procura por ponto seguido de espaço
                period_pos = text.find('. ', search_start, search_end)
                if period_pos != -1:
                    end = period_pos + 1
                else:
                    # Procura por nova linha
                    newline_pos = text.find('\n', search_start, search_end)
                    if newline_pos != -1:
                        end = newline_pos
                    else:
                        # Procura por espaço
                        space_pos = text.rfind(' ', start, search_end)
                        if space_pos > start:
                            end = space_pos
            
            # Extrai o chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    'index': chunk_index,
                    'text': chunk_text,
                    'size': len(chunk_text),
                    'start_pos': start,
                    'end_pos': end
                })
                chunk_index += 1
            
            # Move para o próximo chunk com sobreposição
            start = end - self.chunk_overlap
            
            # Evita loop infinito
            if start <= 0 and chunk_index > 0:
                break
        
        return chunks
    
    def process_document(self, content: bytes, filename: str, 
                        mime_type: str = None) -> Dict[str, Any]:
        """
        Processa documento completo: extração de texto e chunking
        
        Args:
            content: Conteúdo do arquivo em bytes
            filename: Nome do arquivo
            mime_type: MIME type do arquivo
            
        Returns:
            Dicionário com texto completo e chunks
        """
        # Valida arquivo
        if not self.is_supported_file(filename, mime_type):
            raise ValueError(f"Tipo de arquivo não suportado: {filename}")
        
        # Extrai texto
        print(f"[doc_processor] Extraindo texto de {filename}...")
        text = self.extract_text(content, filename, mime_type)
        
        if not text or not text.strip():
            raise ValueError("Não foi possível extrair texto do documento")
        
        # Cria chunks
        print(f"[doc_processor] Criando chunks (size={self.chunk_size}, overlap={self.chunk_overlap})...")
        chunks = self.create_chunks(text)
        
        # Calcula hash do conteúdo
        content_hash = self.calculate_hash(content)
        
        return {
            'text': text,
            'chunks': chunks,
            'content_hash': content_hash,
            'text_length': len(text),
            'chunks_count': len(chunks)
        }


def create_document_processor(chunk_size: int = None, chunk_overlap: int = None) -> DocumentProcessor:
    """
    Factory function para criar um DocumentProcessor
    
    Args:
        chunk_size: Tamanho dos chunks (padrão: 500)
        chunk_overlap: Sobreposição entre chunks (padrão: 50)
        
    Returns:
        Instância de DocumentProcessor
    """
    chunk_size = chunk_size or int(os.environ.get("CHUNK_SIZE", "500"))
    chunk_overlap = chunk_overlap or int(os.environ.get("CHUNK_OVERLAP", "50"))
    
    return DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
